import os
from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
    return True

def fix_cessao(doc):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, "Devedor: Devedor:", "Devedor:")
        
        old_cessionario = "LEDA MARIA SOARES JANOT, brasileira, advogada, casada, portadora do CPF: 021.159.805-49 e RG nº 483293 SSP BA, residente a SMPW Q8 conjunto 3, casa 1, Park way, Brasília/DF, CEP: 71740-803"
        new_cessionario = "{{nome_cessionario}}, {{nacionalidade_cessionario}}, {{profissao_cessionario}}, {{estado_civil_cessionario}}, portadora do CPF: {{cpf_cessionario}} e RG nº {{rg_cessionario}}, residente a {{endereco_cessionario}}, CEP: {{cep_cessionario}}"
        replace_text_in_paragraph(p, old_cessionario, new_cessionario)
        
        replace_text_in_paragraph(p, "5314729-14.2025.8.09.0051", "{{numero_processo}}")
    
    # Check tables for PIX
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "PIX" in cell.text and "{{" not in cell.text:
                    for p in cell.paragraphs:
                        p.add_run(" {{pix}}")

def fix_procuracao(doc):
    replacements = {
        "[Nome]": "{{nome_cedente}}",
        "[Nacionalidade]": "{{nacionalidade_cedente}}",
        "[Estado Civil]": "{{estado_civil_cedente}}",
        "[Profissão]": "{{profissao_cedente}}",
        "[RG]": "{{rg_cedente}}",
        "[CPF]": "{{cpf_cedente}}",
        "[Data Nasc]": "{{data_nascimento}}",
        "[Endereço]": "{{endereco_cedente}}",
        "[CEP]": "{{cep_cedente}}",
        "DR. FÁBIO BATISTA BASTOS": "{{outorgado_nome_1}}",
        "brasileiro, solteiro, advogado": "{{outorgado_nacionalidade_1}}, {{outorgado_estado_civil_1}}, {{outorgado_profissao_1}}",
        "sob o nº 40.115": "sob o nº {{outorgado_oab_1}}",
        "DRA. NATANE ALINE DE CARVALHO MONTEIRO": "{{outorgado_nome_2}}",
        "brasileira, advogada": "{{outorgado_nacionalidade_2}}, {{outorgado_profissao_2}}",
        "OAB/DF nº 63.726": "OAB/DF nº {{outorgado_oab_2}}",
        "CPF n°115.617.116-40": "CPF nº {{outorgado_cpf_2}}",
        "[Número Processo]": "{{numero_processo}}",
        "[Data]": "{{data_contrato}}"
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            replace_text_in_paragraph(p, old, new)

def fix_ciencia(doc):
    # This needs count-based replacement for some fields
    nome_count = 0
    profissao_count = 0
    estado_civil_count = 0
    rg_count = 0
    cpf_count = 0
    endereco_count = 0
    cep_count = 0
    valor_count = 0
    
    for p in doc.paragraphs:
        if "[Nome]" in p.text:
            nome_count += 1
            replace_text_in_paragraph(p, "[Nome]", "{{nome_cedente}}" if nome_count == 1 else "{{nome_cessionario}}")
        if "[Profissão]" in p.text:
            profissao_count += 1
            replace_text_in_paragraph(p, "[Profissão]", "{{profissao_cedente}}" if profissao_count == 1 else "{{profissao_cessionario}}")
        if "[Estado Civil]" in p.text:
            estado_civil_count += 1
            replace_text_in_paragraph(p, "[Estado Civil]", "{{estado_civil_cedente}}" if estado_civil_count == 1 else "{{estado_civil_cessionario}}")
        if "[RG]" in p.text:
            rg_count += 1
            replace_text_in_paragraph(p, "[RG]", "{{rg_cedente}}" if rg_count == 1 else "{{rg_cessionario}}")
        if "[CPF]" in p.text:
            cpf_count += 1
            replace_text_in_paragraph(p, "[CPF]", "{{cpf_cedente}}" if cpf_count == 1 else "{{cpf_cessionario}}")
        if "[Endereço]" in p.text:
            endereco_count += 1
            replace_text_in_paragraph(p, "[Endereço]", "{{endereco_cedente}}" if endereco_count == 1 else "{{endereco_cessionario}}")
        if "[CEP]" in p.text:
            cep_count += 1
            replace_text_in_paragraph(p, "[CEP]", "{{cep_cedente}}" if cep_count == 1 else "{{cep_cessionario}}")
        if "[Valor]" in p.text:
            valor_count += 1
            replace_text_in_paragraph(p, "[Valor]", "{{valor_bruto}}" if valor_count == 1 else "{{valor_liquido}}")
            
        replace_text_in_paragraph(p, "[nacionalidade]", "{{nacionalidade_cedente}}")
        replace_text_in_paragraph(p, "[Data]", "{{data_nascimento}}") # Simplified for now
        replace_text_in_paragraph(p, "[Número]", "{{numero_processo}}")
        replace_text_in_paragraph(p, "[Nome Advogado]", "{{nome_advogado}}")
        replace_text_in_paragraph(p, "[Agência]", "{{agencia}}")
        replace_text_in_paragraph(p, "[Conta]", "{{conta}}")
        replace_text_in_paragraph(p, "[Banco]", "{{banco}}")
        replace_text_in_paragraph(p, "[Nome Cedente]", "{{nome_cedente}}")

def fix_quitacao(doc):
    nome_count = 0
    for p in doc.paragraphs:
        if "[Nome]" in p.text:
            nome_count += 1
            replace_text_in_paragraph(p, "[Nome]", "{{nome_cedente}}")
        
        replacements = {
            "[Nacionalidade]": "{{nacionalidade_cedente}}",
            "[Estado Civil]": "{{estado_civil_cedente}}",
            "[Profissão]": "{{profissao_cedente}}",
            "[RG]": "{{rg_cedente}}",
            "[CPF]": "{{cpf_cedente}}",
            "[Data]": "{{data_nascimento}}",
            "[Endereço]": "{{endereco_cedente}}",
            "[CEP]": "{{cep_cedente}}",
            "[DD/MM/AAAA]": "{{data_negociacao}}",
            "[Número Processo]": "{{numero_processo}}",
            "[Vara/Unidade]": "{{vara_unidade}}",
            "[Comarca]": "{{comarca}}",
            "[Número Processo Origem]": "{{numero_processo_origem}}",
            "[Local]": "{{local}}"
        }
        for old, new in replacements.items():
            replace_text_in_paragraph(p, old, new)

templates = {
    "templates/template-cessao-rpv.docx": fix_cessao,
    "templates/template-procuracao-adjudicia.docx": fix_procuracao,
    "templates/template-dec-ciencia-concord.docx": fix_ciencia,
    "templates/template-dec-quitacao.docx": fix_quitacao
}

for path, func in templates.items():
    if os.path.exists(path):
        doc = Document(path)
        func(doc)
        doc.save(path)
        print(f"Fixed {path}")

# Scan for remaining brackets
for path in templates.keys():
    if os.path.exists(path):
        doc = Document(path)
        for i, p in enumerate(doc.paragraphs):
            if "[" in p.text or "]" in p.text:
                print(f"Remaining bracket in {path} par {i}: {p.text}")
